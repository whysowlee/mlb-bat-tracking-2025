"""표 캡션 재정렬 + 신규 [64] 표 캡션 추가."""
import re, copy
from docx import Document
from docx.text.paragraph import Paragraph

PATH = "_work.docx"
doc = Document(PATH)

CAP_RE = re.compile(r"^표\s*\d+\.$")

def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(prefix)

# 1) 기존 캡션 단락 하나를 템플릿으로 복제하여, [64] 표 앞에 캡션 삽입
template = None
for p in doc.paragraphs:
    if CAP_RE.match(p.text.strip()):
        template = p
        break
ref = find("X_advanced 최종 변수 목록")
new_cap_el = copy.deepcopy(template._p)
ref._p.addnext(new_cap_el)
new_cap = Paragraph(new_cap_el, ref._parent)
# 텍스트만 교체 (서식 보존)
for r in list(new_cap.runs):
    r._element.getparent().remove(r._element)
run = new_cap.add_run("표 0.")
# 템플릿 run 서식 흉내 (italic)
if template.runs:
    run.italic = template.runs[0].italic
    run.bold = template.runs[0].bold

# 2) 문서 순서대로 모든 표 캡션 재번호
n = 0
for p in doc.paragraphs:
    if CAP_RE.match(p.text.strip()):
        n += 1
        # 첫 run 텍스트만 갱신, 나머지 run 제거
        runs = p.runs
        if runs:
            runs[0].text = f"표 {n}."
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
print("총 표 캡션 수:", n)

doc.save(PATH)
print("저장:", PATH)
