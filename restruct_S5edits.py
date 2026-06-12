# -*- coding: utf-8 -*-
"""[38] reliability 작도법 · [57] 과적합/OOF 설명 쉽게 녹이기 · [61][62] Isotonic Mapping 문장화."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix): return p
    return None
def set_text(p, text):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def insert_before(ref, text):
    np = OxmlElement("w:p"); ref._p.addprevious(np); p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    p.add_run(text); return p

ok = {}

# [38] Reliability Diagram 작도법 (해당 단락 앞에 삽입)
r = find("Reliability Diagram에서 선형 모델")
if r:
    insert_before(r,
        "Reliability Diagram(신뢰도 곡선)은 다음과 같이 그린다. ① 예측 확률을 기준으로 타구를 15개의 동일 크기 구간으로 "
        "나누고, ② 각 구간에서 평균 예측 확률(x축)과 실제 안타 비율(y축)을 구해 점을 찍은 뒤, ③ 이 점들이 완벽히 보정된 "
        "상태를 뜻하는 y=x 대각선에 얼마나 가까운지를 본다."); ok["38"]=True

# [57] Isotonic 설명 — 쉬운 말 + OOF(과적합 아님) 녹이기
p1 = find("이 선택은 단순한 경험적 결정이 아니라")
if p1:
    set_text(p1,
        "이 선택은 단순한 경험적 결정이 아니라 본 연구의 목표(Brier 최소화)와 수학적으로 맞아떨어지는 필연적 선택이다. "
        "Isotonic Regression은 예측 확률의 순서(순위)는 그대로 둔 채, 실제 정답(0/1)과의 평균 제곱 오차가 최소가 되도록 "
        "확률 값만 단조적으로 다시 맞추는 보정이다."); ok["57a"]=True
p2 = find("이 목적함수")
if p2:
    set_text(p2,
        "이때 최소화하는 평균 제곱 오차가 바로 Brier Score의 정의식(평균 (y−p)²)과 동일하다. 즉 Isotonic을 적용하는 것 "
        "자체가 Brier를 직접 줄이는 행위이므로, ‘확률 정상도(calibration)를 높이겠다’는 본 연구의 목표와 정확히 일치한다. "
        "다만 같은 데이터에 맞추고 같은 데이터에서 측정하면 과적합이 될 수 있으므로, 본 연구는 Isotonic을 5-fold OOF"
        "(out-of-fold) 방식으로 적합·평가했다. 각 보정 확률이 그 보정을 학습할 때 쓰지 않은 데이터에서 산출되도록 한 것이라, "
        "보고된 Brier 개선은 자기만족이 아니라 held-out 데이터에서 일반화된 결과다. 또한 Isotonic은 순서를 보존하는 단조 "
        "변환이므로 AUC가 대변하는 변수 간 정렬은 전혀 훼손하지 않는다."); ok["57b"]=True

# [61] Isotonic Mapping 왼쪽 문장화
l = find("왼쪽: Stacking raw proba")
if l:
    set_text(l,
        "그림 20의 왼쪽은 모델이 원래 출력한 확률(raw)을 Isotonic이 보정한 확률로 바꿔 주는 변환 곡선이다. 이 곡선은 "
        "순서를 뒤집지 않는 단조 비모수 함수이며, y=x 대각선에서 벗어난 정도가 클수록 그 구간에서 보정이 크게 일어났음을 뜻한다."); ok["61"]=True

# [62] Isotonic Mapping 오른쪽 문장화 (더 자세히)
rt = find("오른쪽: Raw 분포는 0~0.5")
if rt:
    set_text(rt,
        "오른쪽은 보정 전후의 예측 확률 분포다. 보정 전(raw)에는 확률이 0~0.5 구간에 과도하게 몰려 있어 서로 다른 타구들이 "
        "비슷한 확률값으로 뭉뚱그려진다. 반면 Isotonic 보정 후에는 확률이 0 근처(거의 확실한 아웃)와 1 근처(거의 확실한 안타) "
        "양 극단으로 더 넓게 퍼진다. 즉 확실한 타구와 애매한 타구를 더 또렷하게 구분하게 되어 확률의 해상도가 높아지며, 이는 "
        "타구별 확률을 시즌 평균하여 ca-xBA를 산출할 때 선수 간 변별력을 키운다."); ok["62"]=True

doc.save("_w2.docx")
print("S5edits 완료:", ok)
