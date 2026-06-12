"""Word(.docx) 본문 직접 편집 — EDA 섹션 이전까지 코멘트 반영.
칼럼 너비/직접 편집 보존을 위해 in-place 편집. 작업용 _work.docx 대상.
"""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PATH = "_work.docx"
doc = Document(PATH)

def paras():
    return doc.paragraphs

def find(prefix, contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith(prefix) and (contains is None or contains in t):
            return p
    raise LookupError(f"para not found: {prefix!r}")

def find_all(prefix):
    return [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]

def clear_runs(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

def add_rich(p, text):
    # **bold** 마크업 지원
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        r = p.add_run(seg)
        if i % 2 == 1:
            r.bold = True

def set_text(p, text):
    clear_runs(p)
    add_rich(p, text)

def insert_after(ref, text, copy_style_from=None):
    new_p = OxmlElement("w:p")
    ref._p.addnext(new_p)
    p = Paragraph(new_p, ref._parent)
    src = copy_style_from if copy_style_from is not None else ref
    try:
        p.style = src.style
    except Exception:
        pass
    add_rich(p, text)
    return p

def delete_para(p):
    p._element.getparent().remove(p._element)

def delete_table_by_header(cells_sig):
    for tbl in doc.tables:
        try:
            row0 = tbl.rows[0].cells
            sig = [row0[i].text.strip() for i in range(min(len(cells_sig), len(row0)))]
        except Exception:
            continue
        if sig == cells_sig:
            tbl._element.getparent().remove(tbl._element)
            return True
    return False

# ---------------------------------------------------------------------------
# [8] 로드맵 intro — 웅장한 어휘 완화
p = find("본 로드맵은 데이터 누수")
set_text(p, p.text.replace("완벽하게 차단하고", "원천적으로 방지하고"))

# ---------------------------------------------------------------------------
# [20] ballparks intro — Phase1 참조 대신 Athletics 이슈 간단 설명
p = find("MLB 30 개 구장 각각의 물리적 특성")
set_text(p,
    "MLB 30 개 구장 각각의 물리적 특성을 정리한 정적 테이블이다. 행 단위는 구장 1개당 1행"
    "(home_team abbreviation 기준 29 행)이다. Athletics는 2024년 홈구장이 오클랜드 콜리세움,"
    " 2025년이 새크라멘토 서터 헬스 파크로 바뀌어 펜스 거리·고도 등 환경 변수를 단일 구장으로"
    " 매핑할 수 없어 분석에서 제외했기 때문에, 30개 구장 중 29개 행만 남는다. 주요 컬럼은 다음과 같다."
)

# ---------------------------------------------------------------------------
# [29]+[31] attrition 요약 확장 (BIP 필터/도메인 컷오프 상세 + popup 근거 흡수)
p = find("원본 1,443,801행은 5단계")
set_text(p, "원본 1,443,801행은 5단계 정제를 거쳐 최종 225,414개의 BIP로 압축된다(단계별 행 수 변화는 그림 1 Attrition Funnel에 시각화).")
p2 = insert_after(p,
    "가장 큰 두 컷은 (1) BIP 필터와 (2) 도메인 컷오프다. **BIP(Balls In Play, 인플레이 타구) 필터**는 "
    "타자가 친 공이 실제로 그라운드에 들어간 타구(`bb_type`이 ground_ball·fly_ball·line_drive·popup 중 하나)만 "
    "남기고, 삼진·볼넷·헛스윙처럼 타구가 발생하지 않은 투구 행을 제거한다. xBA는 정의상 ‘인플레이 타구의 안타 확률’을 "
    "다루는 지표이므로 타구가 없는 행은 분석 대상이 될 수 없어 이 필터가 필수이며, 이 단계에서 전체의 약 82.5%가 제거된다.")
p3 = insert_after(p2,
    "**도메인 컷오프(|launch_angle| > 60°)**는 발사각이 극단적으로 높거나 낮은 타구를 노이즈로 보고 제거한다(추가 16,248행). "
    "이는 popup(팝업)의 평균 발사각이 65.8°·안타율 1.4%로 사실상 자동 아웃에 가까운 군집이기 때문이며(그림 2에서 popup이 "
    "컷오프로 거의 전량 제거됨을 확인), 이런 극단 타구를 남기면 모델이 학습할 ‘안타 가능 구간’의 신호가 희석된다. 두 기준 모두 "
    "MLB 공식 xBA의 입력 정의와 정렬하기 위한 보수적 선택이다.")

# [31] bb_type 분포 섹션 삭제 (heading + 코드블록)
delete_para(find("BIP 필터 직후 bb_type 분포"))
delete_para(find("bb_type", contains="1190620"))

# ---------------------------------------------------------------------------
# [33] 배트 트래킹 결측률 — 줄글 요약 + 표 삭제
h = find("배트 트래킹 결측률")
summary = insert_after(h,
    "배트 트래킹 변수(bat_speed, swing_length 등)는 2024년 약 11.4%, 2025년 약 5.2%가 결측인데, 이는 해당 트래킹이 "
    "2024 시즌 초반에 비공개였던 구간 때문이다. 결측 행을 제거하면 다량의 데이터를 잃으므로, 트리 모델의 NaN-native 처리를 "
    "활용해 결측치를 그대로 두고 `*_is_missing` 플래그를 추가하여 결측 패턴 자체를 신호로 보존했다.",
    copy_style_from=find("원본 1,443,801행은 5단계"))
delete_para(h)
delete_para(find("표 4.", contains="표 4."))
delete_table_by_header(["", "2024", "2025"])

# ---------------------------------------------------------------------------
# [35] 돔 마스킹 표준값 근거 설명 추가
ref = find("도메인 의의: 트리 앙상블")
insert_after(ref,
    "실내 표준값의 근거는 다음과 같다. 돔(개폐형 포함)이 닫힌 경기는 공조 시스템으로 온·습도가 일정하게 유지되므로, MLB 돔 경기의 "
    "표준 실내 온도인 22°C와 실내 공조의 권장 상대습도 중간값(ASHRAE 기준 50%)을 대표값으로 사용했다. 기압은 실내·외가 동일해 "
    "원본을 유지한다. 외부 기상 5종(풍속·돌풍·풍향·강수·운량)을 0으로 둔 것은 닫힌 돔에서는 외부 기상이 타구에 영향을 줄 수 없다는 "
    "도메인 사실을 데이터에 직접 반영한 것이다.",
    copy_style_from=ref)

# ---------------------------------------------------------------------------
# [38] 기상 변수 요약 통계 — 완전 삭제
delete_para(find("기상 변수 요약 통계"))
delete_para(find("표 5.", contains="표 5."))
delete_table_by_header(["", "count", "mean"])

# ---------------------------------------------------------------------------
# [44] 강수 추측성 문장 삭제(완화)
p = find("강수는 강한 영")
set_text(p, "강수는 0(무강수)에 강하게 집중된 분포를 보인다.")

# ---------------------------------------------------------------------------
# [45] 환경 히트맵 목적 설명 보강
p = find("환경 변수들이 구장 간에 의미 있는 분산")
set_text(p,
    "이처럼 환경 변수(고도·기온·풍속·HR park effect 등)는 구장 간에 뚜렷하고 의미 있는 분산을 가진다. 예컨대 COL(쿠어스 필드)은 "
    "해발 5,190ft 고도로 단일 변수만으로도 분리되고, 돔 구장들은 외부 기상이 차단되어 다른 패턴을 보인다. 이 그림이 필요한 이유는, "
    "바로 이 구장 간 환경 분산이 ca-xBA가 공식 xBA를 넘어 추가로 추출하려는 ‘환경 신호’의 원천임을 보여주기 때문이다 — 환경 변수에 "
    "구장 간 변별력이 없다면 ca-xBA의 개선 여지 자체가 존재하지 않는다.")

# ---------------------------------------------------------------------------
# [50] NaN median 이유 보강
p = find("총 13개 numeric 컬럼의 결측치")
set_text(p,
    "총 13개 numeric 컬럼의 결측치는 2024 전체의 중앙값(median)으로 대체했다. 평균(mean) 대신 중앙값을 쓴 이유는 발사속도·스핀 등 "
    "변수에 이상치가 있어 평균이 왜곡되기 쉽기 때문이고, 2025가 아닌 2024 데이터만으로 대체값을 계산한 이유는 검증 셋(2025)의 정보가 "
    "학습에 새는 데이터 누수를 차단하기 위해서다. 동시에 각 변수의 `*_is_missing` 플래그를 추가해 결측 패턴 자체를 모델 입력 신호로 "
    "보존한다(세부 대체 값은 부록 B 참조).")

# ---------------------------------------------------------------------------
# [52] CV 방법 선정 이유 설명
ref = find("2025는 Phase 5 외부 검증 전용")
insert_after(ref,
    "별도의 80/20 hold-out 시험셋을 두지 않고 5-fold CV를 채택한 이유는, 데이터를 5번 번갈아 학습·검증에 사용함으로써 한정된 2024 "
    "데이터의 활용도를 최대화하고 단일 분할에서 오는 우연한 성능 편차를 줄이기 위함이다. StratifiedKFold는 각 fold의 안타/아웃 비율을 "
    "전체와 동일하게 유지해 평가의 안정성을 높인다.",
    copy_style_from=ref)

# ---------------------------------------------------------------------------
# [54] 제거 변수 — 쌍→제거 변수 형식
p = find("제거된 변수 전체 목록")
set_text(p,
    "고상관 쌍에서 제거한 변수는 다음과 같다(쌍 → 제거): `release_speed` ↔ `effective_speed` 에서 파생 변수인 `effective_speed` 를, "
    "`elevation` ↔ `wx_surface_pressure` 에서 분산이 더 작은 `wx_surface_pressure` 를 제거했고, 동일한 결측 패턴을 공유하는 "
    "배트 트래킹 `*_is_missing` 플래그 7종(예: `bat_speed_is_missing` ↔ `swing_length_is_missing` 등)에서 중복을 제거했다. "
    "전체 24개 고상관 쌍과 각 쌍의 제거 변수는 부록 A에 정리했다.")

# ---------------------------------------------------------------------------
# [56] 스케일링 / Robust Scaler 설명
ref = find("스케일 적용 컬럼: 50개")
insert_after(ref,
    "스케일링은 변수마다 단위와 범위가 제각각인 문제(예: 발사속도 60~120 mph vs 발사각 −60~60°)를 해소해, 거리·계수 기반 연산이 특정 "
    "변수에 치우치지 않도록 만드는 과정이다. 특히 본 데이터에는 이상치가 존재하므로, 평균·표준편차 대신 중앙값과 사분위범위(IQR)로 스케일을 "
    "맞추는 RobustScaler를 사용해 이상치의 영향을 최소화했다(이진 0/1 변수는 제외).",
    copy_style_from=ref)

# ---------------------------------------------------------------------------
# [57] 저장 경로 문구 삭제
p = find("2024 전체 fit, transform")
set_text(p, "스케일러는 2024 전체 데이터로 fit한 뒤 transform을 적용한다.")
delete_para(find("아래 PNG는 모두 pipeline/figures"))

# ---------------------------------------------------------------------------
# [60] RF 튜닝 과정 줄글 설명
ref = find("best CV neg_brier_score")
insert_after(ref,
    "정리하면, Feature Selection의 중요도 계산에 쓰는 Random Forest 자체도 하이퍼파라미터에 민감하므로, RandomizedSearchCV로 "
    "트리 개수·깊이·분할 기준 등을 무작위로 20회 탐색하며 3-fold 내부 CV의 Brier Score가 가장 낮은 조합을 먼저 찾는다. 이렇게 튜닝된 "
    "모델의 `feature_importances_`(분할 시 불순도 감소 기여도)를 변수 중요도로 사용해 신뢰도를 높였다.",
    copy_style_from=ref)

# ---------------------------------------------------------------------------
# [61] 구현 각주 삭제
delete_para(find("비고: 당초 3-model Permutation"))

# ---------------------------------------------------------------------------
# [64] 최종 변수 목록 → 출처 데이터셋별 표
ref = find("X_advanced 최종 변수 목록")
listp = find("launch_speed, launch_angle, bat_speed")
delete_para(listp)
# 표 삽입 (ref 다음)
rows = [
    ("출처 데이터셋", "변수 (개수)"),
    ("Statcast — 타구·배트 트래킹", "launch_speed, launch_angle, bat_speed, swing_length, attack_angle, attack_direction, swing_path_tilt, intercept_ball_minus_batter_pos_x/y_inches, intercept_ball_minus_batter_pos_x_inches_is_missing (10)"),
    ("Statcast — 투구 물리", "release_speed, release_pos_x/z, pfx_x/z, plate_x/z, release_spin_rate, release_extension, spin_axis, api_break_z_with_gravity, api_break_x_arm, api_break_x_batter_in, arm_angle (14)"),
    ("Statcast — 타석 상황", "balls, strikes, outs_when_up, inning, age_pit, age_bat, n_thruorder_pitcher, n_priorpa_thisgame_player_at_bat (8)"),
    ("Statcast — 카테고리(인코딩)", "stand_R, p_throws_R, pitch_type_EP/FC/SC/SL/ST, if_fielding_alignment_(Infield shade/Standard/UNK), of_fielding_alignment_(Standard/Strategic) (12)"),
    ("구장 스펙 (ballparks.csv)", "left_field, center_field, right_field, min_wall_height, max_wall_height, hr_park_effects, extra_distance, elevation, roof, daytime (10)"),
    ("기상 (Open-Meteo)", "wx_temperature_2m, wx_relative_humidity_2m, wx_wind_speed_10m, wx_wind_direction_10m, wx_precipitation, wx_cloud_cover, wx_wind_gusts_10m (7)"),
]
tbl = doc.add_table(rows=len(rows), cols=2)
try:
    tbl.style = doc.tables[0].style
except Exception:
    pass
for i, (a, b) in enumerate(rows):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
# 표를 ref 단락 뒤로 이동
ref._p.addnext(tbl._tbl)

doc.save(PATH)
print("편집 저장 완료:", PATH)
print("표 개수:", len(doc.tables))
