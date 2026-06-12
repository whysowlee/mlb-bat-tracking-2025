# -*- coding: utf-8 -*-
"""신규 코멘트 4건: [29] 샘플링 method 설명 · [30] 재플롯 정합 · [44][45] 문장화."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix, style=None):
    for p in doc.paragraphs:
        if style is not None and p.style.name != style: continue
        if p.text.strip().startswith(prefix):
            return p
    return None
def clear(p):
    for r in list(p.runs): r._element.getparent().remove(r._element)
def set_text(p, text):
    clear(p)
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def insert_after(ref, text):
    np = OxmlElement("w:p"); ref._p.addnext(np)
    p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    p.add_run(text); return p

# [29] 샘플링 비교 헤딩(Heading 3, TOC 아님) 뒤에 method 설명 삽입
samp = find("샘플링 비교 (3종", style="Heading 3")
insert_after(samp,
    "비교한 세 기법은 None(원본 분포를 그대로 사용), RandomUnderSampler(다수 클래스인 아웃을 무작위로 줄여 균형을 "
    "맞춤), SMOTE(소수 클래스인 안타 샘플들 사이를 보간해 가상 샘플을 합성)다. 세 기법 모두 XGBoost default와 동일한 "
    "5-fold CV 위에서 OOF Brier를 비교하여 최종 샘플링을 선정했다.")

# [30] 재플롯(모든 지표 높을수록 우수, Brier·LogLoss=1−값) 정합 텍스트
p30 = find("Brier·LogLoss는 낮을수록")
if p30 is not None:
    set_text(p30,
        "그림 12는 비교 편의를 위해 모든 지표를 ‘높을수록 우수’하도록 통일해 표시했다(원래 낮을수록 좋은 Brier·LogLoss는 "
        "1−값으로 변환). 따라서 막대가 가장 높은 None이 확률 정상도(calibration) 기준 최우수다.")

# [44] XGB 데이터 효과 문장화
p44 = find("데이터 효과 (XGB 위)")
if p44 is not None:
    set_text(p44,
        "반면 비선형 모델(XGBoost) 위에서는 같은 환경 변수를 추가했을 때 Brier가 −0.00423만큼 명확히 개선된다. 트리는 "
        "입력 공간을 국소적으로 분할(split)하며 변수들의 조건부 결합을 학습할 수 있어, 환경 변수의 가치가 비로소 발현되기 때문이다.")

# [45] Interaction 문장화
p45 = find("이 두 값의 차이 = Interaction")
if p45 is not None:
    set_text(p45,
        "선형 모델에서의 데이터 효과(−0.00096)와 비선형 모델에서의 데이터 효과(−0.00423)의 차이가 바로 상호작용"
        "(Interaction = (M4−M2)−(M3−M1) = −0.00327, Brier 감소 방향)이다. 즉 환경 변수의 효과 크기가 알고리즘에 따라 "
        "달라진다는 뜻이며, 이것이 데이터와 알고리즘 간 비선형 상호작용의 정량적 증거다.")

doc.save("_w2.docx")
print("S3edits 완료:",
      "[29]", samp is not None, "[30]", p30 is not None, "[44]", p44 is not None, "[45]", p45 is not None)
