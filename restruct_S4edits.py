# -*- coding: utf-8 -*-
"""4장 코멘트 편집: [50] 소제목 · [51] 모델 설명 · [52] 왜 RandomizedSearchCV · [54] 표13 설명 · [59][64] 문장화."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix, style=None):
    for p in doc.paragraphs:
        if style is not None and p.style.name != style: continue
        if p.text.strip().startswith(prefix): return p
    return None
def clear(p):
    for r in list(p.runs): r._element.getparent().remove(r._element)
def set_text(p, text):
    clear(p)
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def _new_after(ref):
    np = OxmlElement("w:p"); ref._p.addnext(np); return Paragraph(np, ref._parent)
def _new_before(ref):
    np = OxmlElement("w:p"); ref._p.addprevious(np); return Paragraph(np, ref._parent)
def insert_after(ref, text):
    p = _new_after(ref)
    try: p.style = ref.style
    except Exception: pass
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
    return p
def insert_before(ref, text):
    p = _new_before(ref)
    try: p.style = ref.style
    except Exception: pass
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
    return p

ok = {}

# [50] 소제목 변경 (Heading 3)
h = find("Base 모델 튜닝 결과", style="Heading 3")
if h: set_text(h, "Base 모델 구성 및 튜닝 결과"); ok["50"]=True

# [51] (모델 설명) → 모델 소개
m = find("(모델 설명)")
if m:
    set_text(m,
        "본 단계는 서로 다른 트리 앙상블 3종을 base 모델로 사용한다. **Random Forest(RF)**는 여러 결정트리를 독립적으로 "
        "학습해 평균 내는 배깅(bagging) 방식이고, **XGBoost**와 **LightGBM(LGBM)**은 앞선 트리의 오차를 다음 트리가 "
        "보정하며 순차 학습하는 부스팅(boosting) 방식이다. 이 세 모델의 예측을 메타 모델(로지스틱 회귀)로 다시 결합한 것이 "
        "**Stacking**이며, 모델 종류와 무관하게 출력 확률을 실제 빈도에 맞게 단조적으로 재보정하는 비모수 변환이 "
        "**Isotonic Calibration**이다."); ok["51"]=True

# [52] 왜 RandomizedSearchCV
r = find("RF·XGB·LGBM 세 base 모델은")
if r:
    insert_after(r,
        "RandomizedSearchCV를 쓴 이유는, 하이퍼파라미터 후보 공간이 넓을 때 모든 조합을 전수 탐색하는 GridSearch 대신 "
        "무작위로 일부 조합(여기서는 30회)만 추출해 평가함으로써 연산 비용을 크게 줄이면서도 충분히 좋은 조합을 효율적으로 "
        "찾을 수 있기 때문이다."); ok["52"]=True

# [54] 표 13 결과 설명 (6개 후보 단락 앞에 삽입)
s = find("6개 후보 모델의 fold")
if s:
    insert_before(s,
        "표 13은 6개 후보(RF·XGB·LGBM 튜닝, Stacking, Stacking+Isotonic, LGBM+Isotonic)의 OOF 성능이다. 튜닝된 "
        "단일 모델만으로도 Brier가 약 0.131 수준이며, Isotonic 보정을 결합한 LGBM+Isotonic(0.13092)과 "
        "Stacking+Isotonic(0.13083)이 가장 우수하다."); ok["54"]=True

# [59] 최종 선정 문장화 (+ 다음 불릿 병합 삭제)
f = find("최종 선정 모델: LGBM")
if f:
    set_text(f,
        "최종 선정 모델은 LGBM + Isotonic이다. 가장 우수했던 단일 모델 LGBM에 Isotonic 보정을 결합한 이 모델(Brier "
        "0.13092)은 복잡한 Stacking + Isotonic(0.13083)과의 차이가 ΔBrier = +0.00009로, 오캄의 면도날 임계값 "
        "ε(0.001)보다 작아 fold 변동 수준 내의 통계적 동률이다. 따라서 성능이 사실상 같다면 더 단순한 모델을 택한다는 "
        "원칙에 따라 단일 모델 기반 LGBM + Isotonic을 자동 선정했다.")
    nxt = find("Best_Single(LGBM) + Isotonic 와 Stacking")
    if nxt: nxt._p.getparent().remove(nxt._p)
    ok["59"]=True

# [64] Calibration 문장화
c = find("Stacking + Isotonic 이 대각선에")
if c:
    set_text(c,
        "Reliability Diagram에서는 Stacking + Isotonic 곡선이 대각선(y=x)에 가장 가깝게 붙어 있는데, 이는 그 예측 "
        "확률이 실제 안타 빈도와 가장 잘 일치함을 뜻하며 OOF Brier가 최소인 결과와 부합한다."); ok["64"]=True

doc.save("_w2.docx")
print("S4edits 완료:", ok)
