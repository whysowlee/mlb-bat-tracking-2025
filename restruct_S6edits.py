# -*- coding: utf-8 -*-
"""[66] R²/표16 설명 줄글 · [76] 표 20 'Top N 적중' 칼럼 채우기."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix): return p
    return None
def insert_after(ref, text):
    np = OxmlElement("w:p"); ref._p.addnext(np); p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    p.add_run(text); return p

ok = {}

# [66] R² 사용 이유 + 표16 의미
ref = find("MLB 공식 xBA (est_ba) R²")
if ref:
    insert_after(ref,
        "본 검증에서 R²(결정계수)를 사용한 이유는, ca-xBA가 선수의 실제 타격 성과(wOBA)를 얼마나 잘 설명하는지를 하나의 "
        "숫자로 요약하기 위해서다. R²는 실제 wOBA의 분산 중 해당 지표로 설명되는 비율(0~1)로, 1에 가까울수록 그 지표가 "
        "실제 성과를 정확히 대변한다는 뜻이다. 표 16은 동일한 309명에 대해 ca-xBA와 MLB 공식 xBA를 각각 실제 wOBA와 "
        "1:1로 비교한 것으로, ca-xBA의 R²(0.3976)가 공식 xBA(0.2499)보다 높다. 즉 ca-xBA가 선수의 실제 기량을 "
        "상대적으로 약 59% 더 많이 설명하며, 이것이 본 연구 모델이 공식 지표보다 우수하다는 핵심 정량 근거다.")
    ok["66"]=True

# [76] 표 20 'Top N 적중' 칼럼 채우기 (순위 기준)
for tbl in doc.tables:
    hdr = [c.text.strip() for c in tbl.rows[0].cells]
    if "Top N 적중" in hdr:
        ri = hdr.index("본 연구의 ca-xBA 순위"); ti = hdr.index("Top N 적중")
        filled = 0
        for row in tbl.rows[1:]:
            rank = row.cells[ri].text.strip()
            if rank in ("—", "-", ""):
                val = "검증불가"
            else:
                try: val = "적중" if int(re.sub(r"[^0-9]", "", rank)) <= 10 else "미달"
                except ValueError: val = "검증불가"
            row.cells[ti].text = val; filled += 1
        ok["76"] = f"{filled}행 채움"
        break

doc.save("_w2.docx")
print("S6edits 완료:", ok)
