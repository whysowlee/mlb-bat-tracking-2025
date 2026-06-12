# -*- coding: utf-8 -*-
"""S2: Brier 블록을 FS 앞으로 이동 + 코멘트 줄글 편집 13건 + 표5 수정."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix, exact=False):
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t == prefix) if exact else t.startswith(prefix):
            return p
    return None
def find_all(prefix):
    return [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
def clear(p):
    for r in list(p.runs): r._element.getparent().remove(r._element)
def add_rich(p, text):
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def set_text(p, text):
    clear(p); add_rich(p, text)
def insert_after(ref, text):
    np = OxmlElement("w:p"); ref._p.addnext(np)
    p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    add_rich(p, text); return p

# ── Brier 블록을 FS 앞으로 이동 ──────────────────────────────────────────
brier_h = find("평가 지표 Brier Score 정의")
samp_h = find("샘플링 비교")
fs_h = find("Feature Selection (RF importance")
# brier_h ~ samp_h 직전까지 단락 수집 (본문 문단만; 표 없음)
blk = []
node = brier_h._p
while node is not None and node is not samp_h._p:
    blk.append(node); node = node.getnext()
for node in blk:
    fs_h._p.addprevious(node)
print("Brier 블록 → FS 앞 이동:", len(blk), "요소")

# ── [19] (2) ATH 제외 설명 ───────────────────────────────────────────────
set_text(find("(2) ATH 제외", exact=True) or find("(2) ATH 제외"),
    "(2) ATH 제외: Athletics(오클랜드 애슬레틱스)는 2024년 홈구장이 오클랜드 콜리세움, 2025년이 새크라멘토 서터 "
    "헬스 파크로 달라 펜스 거리·고도 등 환경 변수를 단일 구장으로 매핑할 수 없어, 해당 홈경기 타구를 분석에서 제외했다(약 8,641행).")

# ── [20] (3) 물리 결측 제거 → launch 결측 행 제거 설명 ───────────────────
p3 = find("(3) 물리 결측 제거")
set_text(p3,
    "(3) 물리 결측 제거: launch_speed(발사속도) 또는 launch_angle(발사각)이 NaN인 행을 제거한다(약 2,878행). "
    "xBA 자체가 발사속도 × 발사각의 함수라 이 두 값이 없으면 모델 입력 자체가 불가능하며, 손실 비율이 작아 행 제거 방식을 채택했다.")

# ── [22] 마스킹 통계 풀어쓰기 ─────────────────────────────────────────────
set_text(find("마스킹된 BIP 행:"),
    "지붕을 여닫을 수 있는 구장(retractable)과 완전 돔 구장(TB)에서 나온 타구 61,677개 중에서, 실제로 그날 지붕이 "
    "닫혀 있던 경기의 타구 43,197개(70.0%)에 대해 외부 기상 데이터를 마스킹(상수로 덮어쓰기)했다.")

# ── [26] 분포 동일 줄글화 ─────────────────────────────────────────────────
set_text(find("2024와 2025의 분포가 거의 동일"),
    "2024년과 2025년의 발사속도·발사각 분포가 거의 동일하여, 연도별 분리(Temporal Split) 이후에도 입력 변수의 분포가 "
    "안정적으로 유지됨을 시각적으로 확인할 수 있다.")

# ── [27] 히트맵 대시 제거 문장화 ─────────────────────────────────────────
set_text(find("이 히트맵은 본 연구가 트리"),
    "이 히트맵은 본 연구가 트리 앙상블 알고리즘을 채택한 수학적 근거이기도 하다. 그림 4에서 보듯 발사속도와 발사각이 "
    "안타율에 미치는 영향은 그 자체로 이미 강한 비선형성을 띤다. 특정 각도와 속도가 교차하는 좁은 구간에서만 안타율이 "
    "급증하기 때문이다. 여기에 온도·풍속·구장 고도 등 수십 개의 환경 변수가 더 얽히면, 변수 간 독립성과 단조 증가를 "
    "가정하는 선형 모델로는 이런 복잡한 교호작용을 담아낼 수 없다. 따라서 입력 공간을 조건부로 분할하는 비선형 트리 "
    "모델이 필요하다(Phase 3에서 통계적으로 검정).")

# ── [29] COL 쉬운 말로 ───────────────────────────────────────────────────
set_text(find("COL(쿠어스):"),
    "예를 들어 COL(쿠어스 필드)은 해발 5,190ft의 높은 고도 때문에 공기가 얇아 타구가 멀리 뻗는 극단적인 구장으로, "
    "고도 하나만으로도 다른 구장과 뚜렷이 구분된다.")

# ── [32] 표5 (가변→실제 개수, 82 복원) + 노트 ───────────────────────────
for tbl in doc.tables:
    if tbl.rows[0].cells[0].text.strip() == "그룹":
        for row in tbl.rows:
            c0 = row.cells[0].text.strip()
            if c0.startswith("(d2) pitch_type"): row.cells[1].text = "16"
            elif c0.startswith("(e2) alignment"): row.cells[1].text = "7"
        newr = tbl.add_row()
        newr.cells[0].text = "X_advanced 초기(One-Hot 후)"; newr.cells[1].text = "82"; newr.cells[2].text = ""
        break
set_text(find("표에서 pitch_type과 fielding_alignment"),
    "표의 변수 수는 카테고리형 변수(pitch_type·fielding_alignment)를 one-hot 인코딩으로 모두 펼친 직후를 기준으로 하며, "
    "이때 합계는 82개다. 이 82개가 이후 다중공선성 제거(−9)와 Feature Selection(−12)을 거쳐 최종 61개로 확정된다.")

# ── [38] 샘플링 CM 화살표 문장화 ─────────────────────────────────────────
cm = find("원본(None): True Negative")
set_text(cm,
    "세 샘플링의 OOF 혼동행렬을 비교하면, 원본(None)은 True Negative가 많고 Recall이 낮은 보수적 예측을 보인다. "
    "언더샘플링은 True Positive가 크게 늘지만 동시에 False Positive도 증가해, Recall은 오르고 Precision은 떨어지는 "
    "트레이드오프가 나타난다. SMOTE는 원본에 가까운 균형을 보이며 F1이 원본보다 약간 높다.")
# 뒤따르던 Under / SMOTE 불릿 제거
for pref in ("Under: TP 대폭 증가", "SMOTE: 원본에 가까운 균형"):
    q = find(pref)
    if q is not None: q._p.getparent().remove(q._p)

# ── [39] None 선정 이유 줄글 ─────────────────────────────────────────────
insert_after(find("Brier·LogLoss는 낮을수록"),
    "본 연구의 목표는 단순 분류가 아니라 정확한 확률 산출이다. 샘플링으로 클래스 균형을 맞추면 안타를 더 많이 잡아내 "
    "(Recall↑) F1은 오를 수 있지만, 그 대가로 모델이 출력하는 확률이 실제 빈도에서 멀어져 왜곡된다. ca-xBA는 이 "
    "확률값을 시즌 평균하여 산출하는 지표이므로 확률 정상도(calibration)가 최우선이며, 따라서 OOF Brier가 가장 낮은 "
    "원본(None)을 최종 샘플링으로 선정했다.")

# ── [52] Effect Decomposition 표 해석 보강 ──────────────────────────────
set_text(find("해석 가이드: Brier"),
    "위 표는 2x2 설계에서 데이터 효과와 알고리즘 효과를 분리한 것이다. 각 행은 한 요인만 바꿨을 때의 성능 변화량(Δ)을 "
    "뜻하며, Brier·LogLoss는 값이 줄수록(음수) 좋고 F1·AUC는 늘수록(양수) 좋다. 특히 마지막 Interaction 행은 "
    "‘데이터를 추가했을 때의 효과가 알고리즘에 따라 얼마나 달라지는가’를 나타내며, 이 값이 0에서 유의하게 벗어날수록 "
    "데이터와 알고리즘 사이의 비선형 상호작용이 강하다는 의미다.")

# ── [53] 데이터 효과 줄글 쉽게 ───────────────────────────────────────────
set_text(find("데이터 효과 (LogReg 위)"),
    "선형 모델(LogReg) 위에서 환경 변수 약 60개를 추가했을 때 Brier 개선은 −0.00096에 그친다. 즉 선형 모델은 변수만 "
    "늘려도 거의 나아지지 않는데, 이는 각 변수가 독립적·단조적으로만 작용한다고 가정해 변수 간 상호작용을 활용하지 "
    "못하기 때문이다.")

# ── [59] 소제목 ──────────────────────────────────────────────────────────
set_text(find("(수리적 본질"),
    "선형 vs 트리 모델의 수리적 구조 — 왜 비선형 모델에서만 환경 변수가 발현되는가")

# ── [64] 트리 앙상블 정당성(calibration) 보강 ───────────────────────────
insert_after(find("대각선에 가까울수록 확률 정상도 양호"),
    "이 Reliability Diagram은 트리 앙상블을 써야 하는 핵심 근거를 보여준다. 선형 모델(M1·M3)은 예측 확률이 0.2~0.4 "
    "구간에 몰려 대각선에서 크게 벗어나는데, 이는 안타 확률이 변수들의 비선형 조합으로 결정됨에도 선형 모델이 이를 단조 "
    "결합으로만 근사하기 때문이다. 반면 트리 모델(M2·M4)은 입력 공간을 조건부로 분할해 국소적으로 다른 확률을 학습하므로 "
    "예측 확률이 실제 안타 빈도와 훨씬 잘 일치한다(대각선에 근접). ca-xBA가 확률값 자체를 평균해 산출되는 지표인 만큼, "
    "이러한 확률 정상도의 우위가 곧 트리 앙상블 채택의 학술적 정당성이다.")

doc.save("_w2.docx")
print("S2 완료")
