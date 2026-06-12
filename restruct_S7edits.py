# -*- coding: utf-8 -*-
"""[71] BABIP 소개 · [73][75][77] 보고서 톤으로 문장 수정."""
import re
from docx import Document

doc = Document("_w2.docx")

def find(prefix, exact=False):
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t == prefix) if exact else t.startswith(prefix): return p
    return None
def set_text(p, text):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    p.add_run(text)

ok = {}

# [71] BABIP 소개 ('(내용)' 자리)
m = find("(내용)", exact=True)
if m:
    set_text(m,
        "BABIP(Batting Average on Balls In Play)은 삼진과 홈런을 제외하고 그라운드에 들어간 인플레이 타구만을 대상으로 "
        "산출한 타율로, 야구 통계에서 ‘운’을 가늠하는 가장 대중적인 지표다. 인플레이 타구가 안타로 연결되는지는 타자의 실력뿐 "
        "아니라 수비수의 위치·구장 환경·우연 등 통제하기 어려운 요소에 크게 좌우되기 때문이다. 다만 BABIP이 리그 평균보다 "
        "높다는 사실만으로 ‘운이 좋았다’고 단정할 수는 없다. 타자마다 고유한 BABIP 수준이 다르므로(빠른 발이나 강한 타구를 "
        "가진 타자는 통산 BABIP이 본래 높다), 진정한 운·불운은 시즌 BABIP을 그 선수의 통산 BABIP(개인 기준선)과 비교한 "
        "편차(Δ_BABIP)로 판단한다. 본 분석은 이 도메인 정통 기준을 ca-xBA 기반 luck 지표를 교차 검증하는 잣대로 사용한다.")
    ok["71"]=True

# [73] 톤 수정
p73 = find("두 지표 모두 양의 상관을 보이지만")
if p73:
    set_text(p73,
        "luck은 시즌 BABIP 및 Δ_BABIP 모두와 양의 상관을 보이나, 개인 기준선을 보정한 Δ_BABIP와의 상관이 도메인적으로 더 "
        "타당하다. 본 분석에서 luck과 Δ_BABIP의 Pearson 상관계수는 0.519로, ca-xBA 기반 luck 지표가 야구 도메인의 정통 "
        "행운 신호(통산 BABIP 대비 시즌 편차)와 동일한 방향을 가리킴을 보여준다. 상관계수가 1.0에 미치지 않는 것은 ca-xBA가 "
        "BABIP 단일 지표로는 포착되지 않는 환경 보정 신호(dome × weather 상호작용, hr_park_effects, 구장 펜스 거리 등)를 "
        "추가로 반영하기 때문이며, 이는 §4.6의 Trout·Schwarber 패턴에서 구체적으로 확인된다.")
    ok["73"]=True

# [75] 톤 수정 (운/행운 해석 가이드)
p75 = find("luck (= BIP-AVG − ca-xBA) 가 양수면")
if p75:
    set_text(p75,
        "luck이 양수인 타자는 contact quality 대비 더 많은 안타가 나온 경우다. 이때 Δ_BABIP > 0(자기 통산 대비 시즌 "
        "BABIP이 높음)을 함께 만족하면 두 지표가 모두 행운 효과를 가리키는 이중 검증에 해당하고, Δ_BABIP가 0에 가깝거나 "
        "음수이면 luck이 포착한 행운이 BABIP 단일 지표로는 드러나지 않는 ca-xBA의 환경 보정 신호임을 의미한다.")
    ok["75"]=True

# [77] 톤 수정 (불운 해석 가이드)
p77 = find("luck 가 음수면 contact quality")
if p77:
    set_text(p77,
        "luck이 음수인 타자는 contact quality 대비 안타가 적게 나온 경우다. Δ_BABIP < 0이면 자기 통산 대비 시즌 BABIP도 "
        "낮아 두 지표가 모두 불운을 가리키며, Δ_BABIP가 0에 가깝거나 양수인데 luck만 크게 음수이면 Trout 패턴에 해당한다. "
        "이 경우 ca-xBA는 ‘이 정도 타구 질이면 더 좋은 결과가 나왔어야 한다’고 평가하지만 BABIP만으로는 불운으로 드러나지 "
        "않으므로, Front Office의 저평가 선수 발굴 포인트가 된다.")
    ok["77"]=True

doc.save("_w2.docx")
print("S7edits 완료:", ok)
