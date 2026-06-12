"""Stage 6: 로드맵 Phase2 항목 순서 교체(FS↔샘플링) + FS 문구 보정 + 모델 명확화."""
import re
from docx import Document

doc = Document("_w2.docx")

def add_rich(p, text):
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def set_text(p, text):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    add_rich(p, text)
def append_run(p, text):
    p.add_run(text)

paras = doc.paragraphs

# 로드맵 영역 식별 (단계별 수행 로드맵 ~ 데이터 전처리 및 탐색적 분석)
rm_start = rm_end = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if t.startswith("단계별 수행 로드맵"): rm_start = i
    if t == "데이터 전처리 및 탐색적 분석": rm_end = i; break

# 로드맵 내 샘플링 / FS 리스트 항목 찾기
samp_li = fs_li = None
for p in paras[rm_start:rm_end]:
    t = p.text.strip()
    if t.startswith("원본(None)·언더샘플링·SMOTE"): samp_li = p
    if t.startswith("RF importance와 Mutual Information 2개 기준"): fs_li = p
# FS를 샘플링 앞으로 (현재 샘플링#3, FS#4 → FS#3, 샘플링#4)
if samp_li is not None and fs_li is not None:
    samp_li._p.addprevious(fs_li._p)
    print("로드맵 FS↔샘플링 순서 교체 완료")

# FS 본문 문구 보정 — 샘플링이 뒤로 갔으므로 forward-ref 제거
for p in doc.paragraphs:
    if p.text.strip().startswith("학습 데이터: 최적 샘플링"):
        set_text(p, "학습 데이터: 원본 분포(별도 샘플링 미적용) 위에서 RF importance와 MI를 계산했다. "
                    "이는 뒤의 샘플링 비교(2.5절)에서 원본(None)이 최적으로 확정되는 결과와 일치한다.")
        print("FS 문구 보정 완료")
        break

# 모델 명확화 ① 초록
for p in doc.paragraphs:
    if p.text.strip().startswith("본 연구는 메이저리그(MLB) 공식 기대 타율(xBA)이 타구의 물리적 질"):
        append_run(p, " 본 연구의 학습 과제는 타구 단위의 이진 분류(지도학습)로, 각 인플레이 타구의 안타 여부"
                      "(안타=1 / 아웃=0)를 예측하며, 이렇게 산출된 안타 예측 확률을 선수별로 시즌 평균한 값이 최종 지표 ca-xBA다.")
        print("초록 모델 명확화 완료")
        break

# 모델 명확화 ② 프로젝트 목표
for p in doc.paragraphs:
    if p.text.strip().startswith("본 프로젝트에서는 타구의 물리 데이터에"):
        append_run(p, " 구체적으로 모델은 각 타구의 안타 여부(1/0)를 맞히는 이진 분류기로 학습되며, ca-xBA는 그 분류기가 "
                      "출력한 안타 확률을 선수 단위로 평균한 세이버메트릭스 지표다. 따라서 본 연구에서 중요한 것은 단순 분류 "
                      "정확도가 아니라 예측 ‘확률값의 정밀도(calibration)’이며, 이를 위해 Brier Score를 핵심 평가 지표로 삼는다.")
        print("프로젝트 목표 모델 명확화 완료")
        break

doc.save("_w2.docx")
print("Stage 6 완료")
