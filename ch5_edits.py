# -*- coding: utf-8 -*-
"""5장 수정: Y축 문구 삭제 · ATH 8명 R² 영향 설명 · 표20 포지션내 순위 주석."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix): return p
    return None
def insert_after(ref, text):
    np = OxmlElement("w:p"); ref._p.addnext(np); p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    p.add_run(text); return p

# 1) 'Y축 기준점 ...' 문구 삭제
y = find("Y축 기준점 (실제 기량)")
if y is not None:
    y._p.getparent().remove(y._p); print("Y축 문구 삭제")

# 2) ATH 8명 → R² 영향 제한적 설명 추가
a = find("PA tolerance (50%) 미달")
if a is not None:
    insert_after(a,
        "이처럼 ATH 소속 일부 선수는 타구 표본이 다소 적은 상태로 비교에 포함되어 있다. 다만 (i) 309명 전원이 ID로 정확히 "
        "매칭되었고, (ii) our_bip/csv.bip 평균 비율이 0.90으로 높으며, (iii) 표본 부족 사례가 8명(대부분 동일 원인인 ATH "
        "홈경기 제외)으로 한정·식별되어 있으므로, 이들이 R²(309명 기준)에 미치는 영향은 제한적이라고 판단된다.")
    print("ATH 설명 추가")

# 3) 표 20: 포지션 내 순위 주석 (표 다음에 삽입)
tbl20 = None
for tbl in doc.tables:
    if "Top N 적중" in [c.text.strip() for c in tbl.rows[0].cells]:
        tbl20 = tbl; break
if tbl20 is not None:
    # 표 다음 형제 위치에 주석 단락 삽입
    note = OxmlElement("w:p"); tbl20._tbl.addnext(note)
    # 본문 스타일(직전 표 캡션 스타일 흉내) — 그냥 일반 단락
    pr = OxmlElement("w:pPr"); ps = OxmlElement("w:pStyle"); ps.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val","Compact"); pr.append(ps); note.insert(0,pr)
    para = Paragraph(note, None)
    para.add_run("※ ‘ca-xBA 순위’는 해당 포지션 내에서의 순위이므로, 서로 다른 포지션의 1위(예: AL 포수 Cal Raleigh, "
                 "AL 외야수 Aaron Judge)가 함께 나타날 수 있다. 마지막 wOBA 열은 각 수상자의 실제 시즌 성과(검증 기준 지표)다.")
    print("표 20 주석 추가")

doc.save("_w2.docx")
print("ch5_edits 완료")
