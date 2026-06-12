"""Stage 4: 코멘트 기반 줄글 설명 추가/보강 ([21][28/29][40][41][45][46][47])."""
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

doc = Document("_w2.docx")

def find(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(prefix)

def clear(p):
    for r in list(p.runs): r._element.getparent().remove(r._element)
def add_rich(p, text):
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "": continue
        r = p.add_run(seg); r.bold = (i % 2 == 1)
def set_text(p, text):
    clear(p); add_rich(p, text)
def insert_after(ref, text):
    np = OxmlElement("w:p"); ref._p.addnext(np)
    p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    add_rich(p, text); return p
def insert_before(ref, text):
    np = OxmlElement("w:p"); ref._p.addprevious(np)
    p = Paragraph(np, ref._parent)
    try: p.style = ref.style
    except Exception: pass
    add_rich(p, text); return p

# [21] Temporal Split 안타율 동일 해석
insert_after(find("2024_data 는 Phase 2~4 학습/평가용"),
    "두 시즌의 안타율이 0.3411로 정확히 동일하다는 점에 주목할 만하다. 이는 연도로 데이터를 분리했음에도 "
    "target(안타)의 기저 비율이 이동하지 않았다는 뜻으로, 2024로 학습한 모델을 2025에 적용할 때 클래스 분포 변화"
    "(label shift)로 인한 왜곡 없이 공정하게 검증할 수 있음을 보증한다.")

# [28][29] '가변' 설명 — NaN 처리 헤딩 앞(=변수 그룹 표 뒤)에 삽입
insert_before(find("NaN 처리"),
    "표에서 pitch_type과 fielding_alignment의 변수 수를 ‘(가변)’으로 표기한 이유는, 이들을 one-hot 인코딩하면 "
    "데이터에 실제 등장하는 카테고리 수만큼 더미 컬럼이 생겨 변수 개수가 고정되지 않기 때문이다. 최종 개수는 이후 "
    "다중공선성 제거와 Feature Selection을 거쳐 확정된다.")

# [40] RF 튜닝 파라미터 결정 과정
set_text(find("정리하면, Feature Selection의 중요도 계산"),
    "정리하면, Feature Selection의 중요도 계산에 쓰는 Random Forest 자체도 하이퍼파라미터에 민감하므로, "
    "트리 개수(100·200·500), 최대 깊이(10·20·무제한), 분할 최소 표본 수, 분할 기준(gini·entropy)의 후보 조합을 "
    "RandomizedSearchCV로 무작위 20회 추출하여 3-fold 내부 CV의 Brier Score가 가장 낮은 조합을 선택했다. 그 결과 "
    "n_estimators=200, max_depth=무제한, min_samples_split=4, criterion=entropy가 best params로 선정되었고, "
    "이렇게 튜닝된 모델의 feature_importances_(분할 시 불순도 감소 기여도)를 변수 중요도로 사용했다.")

# [41] FS 제거 규칙 + X_BASE 제외 이유
set_text(find("제거 규칙: RF importance & MI"),
    "제거 규칙은 RF importance와 Mutual Information 두 지표 **모두에서** 하위 30%에 동시 진입한 변수만 제거하는 "
    "보수적 방식을 채택했다. 두 지표 중 하나라도 중요하다고 판단하면 변수를 보존함으로써, 단일 지표의 편향으로 유용한 "
    "변수가 잘못 제거되는 위험을 줄이기 위함이다. 단, launch_speed·launch_angle로 구성된 X_BASE는 MLB 공식 xBA의 "
    "입력이자 본 연구 통제군(Phase 3)의 기준 변수이므로, 중요도 순위와 무관하게 제거 대상에서 항상 제외했다.")

# [45] KDE 상세 설명
insert_after(find("환경 변수(기온·풍속·고도·HR park effects)"),
    "정리하면, KDE는 각 변수의 값 분포를 안타(is_hit=1)/아웃(0) 그룹으로 나눠 겹쳐 그린 것으로, 두 곡선이 많이 "
    "어긋날수록 그 변수 단독으로 안타 여부를 잘 가른다는 뜻이다. 발사속도·발사각은 두 그룹의 곡선이 뚜렷이 분리되어 "
    "강한 단변량 신호를 보이는 반면, 환경 변수들은 곡선이 거의 포개져 단변량 변별력은 약하다. 이는 환경 변수의 가치가 "
    "다른 변수와의 비선형 상호작용에서 비로소 발현됨을 시사하며, 트리 앙상블 도입의 근거가 된다.")

# [46][47] Boxplot — 추가 이상치 제거 불필요 설명
insert_after(find("launch_speed 의 IQR이 is_hit=1"),
    "박스플롯의 수염(whisker) 밖 점들이 일부 보이지만, 이는 실제 경기에서 나오는 정상적인 타구 분포의 꼬리이며, "
    "popup(±60° 컷오프로 이미 제거)을 제외하면 별도의 이상치 제거가 필요하지 않다. 극단값을 인위적으로 잘라내면 약한 "
    "타구나 강한 타구 같은 실제 신호까지 손실되므로, 본 연구는 도메인 컷오프(popup) 외의 추가 이상치 제거를 수행하지 "
    "않았다. 또한 스케일링에 이상치에 강건한 RobustScaler를 사용해 이러한 꼬리값의 영향을 자연스럽게 완화했다.")

doc.save("_w2.docx")
print("Stage 4 완료")
